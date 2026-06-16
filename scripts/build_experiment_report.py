# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_run_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph(paragraph, first_line=True):
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(4)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(21)


def add_p(doc, text="", size=10.5, first_line=True, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    set_paragraph(p, first_line=first_line)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, size=14 if level == 1 else 12, bold=True, color=(24, 104, 219) if level == 1 else None)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run_font(r, size=9.5, bold=bold)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        shade_cell(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    doc.add_paragraph()
    return table


def build_report(output_path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)

    p = doc.add_paragraph("课程编号：A0801050230")
    set_paragraph(p, first_line=False)
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("软件质量保证与测试\n实验报告")
    set_run_font(r, size=22, bold=True)
    doc.add_paragraph("\n")

    table = doc.add_table(rows=7, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cover_rows = [
        ["组长姓名", "刘皓天", "学号", "20245897"],
        ["班级", "软件2408", "指导教师", "王莹"],
        ["开设学期", "2025-2026 春季学期", "开设时间", "第15周--第16周"],
        ["报告日期", "2026.6.16", "评定成绩", ""],
        ["评定人", "王莹", "评定日期", "2026.6.19"],
        ["", "", "", ""],
        ["", "", "", ""],
    ]
    labels = {"组长姓名", "学号", "班级", "指导教师", "开设学期", "开设时间", "报告日期", "评定成绩", "评定人", "评定日期"}
    for row, values in zip(table.rows, cover_rows):
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value, bold=value in labels)
    doc.add_paragraph("\n\n")
    school = doc.add_paragraph("东北大学软件学院")
    school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(school.runs[0], size=14, bold=True)
    doc.add_page_break()

    add_heading(doc, "1. 实验目的")
    add_p(doc, "（1）练习白盒测试用例设计，熟悉 JUnit 单元测试工具，并使用单元测试方法验证后端核心业务逻辑。")
    add_p(doc, "（2）练习黑盒测试用例设计，从用户角度对大学生个人助手系统进行手工功能测试。")
    add_p(doc, "（3）通过测试发现待测程序的可用性问题，并结合测试结果对程序进行修改和完善。")

    add_heading(doc, "2. 测试计划")
    add_heading(doc, "2.1 实验1测试计划", 2)
    add_p(doc, "（1）系统功能概述")
    add_p(doc, "本组开发的待测程序为 JavaWeb 大学生个人助手，目标用户为在校大学生。系统采用 Vue 3 + Spring Boot + MySQL 前后端分离架构，主要功能包括登录注册、首页数据看板、课程与日程管理、学习任务分析、生活健康记录、消费与预算管理、个人目标与习惯打卡。")
    add_p(doc, "（2）测试目标")
    add_p(doc, "实验 1 的测试目标是通过白盒测试验证后端核心业务逻辑的正确性、异常分支处理能力和用户数据隔离能力。重点检查注册登录、课程查询、任务统计、学习分析、预算预警、健康记录更新、习惯打卡和首页聚合等服务层逻辑。")
    add_p(doc, "（3）测试范围")
    add_p(doc, "测试范围包括 AuthService、CourseService、TaskService、StudyRecordService、BudgetService、HealthRecordService、HabitService 和 DashboardService 等后端服务类。前端页面展示和浏览器兼容性不属于实验 1 的主要范围。")
    add_p(doc, "（4）测试环境")
    add_p(doc, "操作系统为 Windows 11，JDK 23，Maven 3.6.1，MySQL 8.0，Spring Boot 3.2.5，JUnit 5，AssertJ，Spring Boot Test。")
    add_p(doc, "（5）测试工具")
    add_p(doc, "使用 JUnit 5 编写和运行单元测试，使用 Maven Surefire 执行测试套件，使用 AssertJ 进行断言。")
    add_p(doc, "（6）测试策略")
    add_p(doc, "采用语句覆盖、判定覆盖、条件覆盖和异常路径测试相结合的方法。对正常输入、错误输入、边界输入、重复数据和跨用户访问场景分别设计测试用例，重点验证分支逻辑是否被充分覆盖。")

    add_heading(doc, "2.2 实验2测试计划", 2)
    add_p(doc, "（1）测试目标")
    add_p(doc, "实验 2 的测试目标是从普通用户角度验证系统主要功能是否可用，确认界面操作、数据保存、统计展示和异常提示是否满足大学生日常使用需要。")
    add_p(doc, "（2）测试范围")
    add_p(doc, "测试范围覆盖注册登录、首页看板、课程管理、任务 DDL、日程管理、学习记录、健康记录、收支记录、预算预警、目标管理、习惯打卡和个人资料维护。")
    add_p(doc, "（3）测试环境")
    add_p(doc, "浏览器为 Microsoft Edge 或 Chrome，前端地址为 http://localhost:5173，后端地址为 http://localhost:8080，数据库为 MySQL 8.0。")
    add_p(doc, "（4）测试工具")
    add_p(doc, "主要使用浏览器进行手工测试，使用 MySQL 客户端检查数据库结果，使用 Maven 和 npm 分别验证后端测试和前端构建。")
    add_p(doc, "（5）测试策略")
    add_p(doc, "采用等价类划分、边界值分析、场景法和错误推测法设计测试用例。测试过程覆盖正常业务流程、必填项为空、重复数据、状态变化和跨模块统计展示。")

    add_heading(doc, "3. 测试用例")
    add_heading(doc, "3.1 实验1测试用例及结果分析", 2)
    add_p(doc, "（1）测试用例")
    white_rows = [
        ["W01", "用户注册", "等价类/异常路径", "注册唯一用户名", "返回用户资料，用户名保存成功", "通过"],
        ["W02", "重复用户名注册", "判定覆盖", "同一用户名注册两次", "抛出“用户名已存在”异常", "通过"],
        ["W03", "错误密码登录", "异常路径", "用户名正确、密码错误", "抛出“用户名或密码错误”异常", "通过"],
        ["W04", "今日课程过滤", "条件覆盖", "创建今日课程和非今日课程", "仅返回今日课程", "通过"],
        ["W05", "课程用户隔离", "路径覆盖", "用户 A 修改用户 B 课程", "返回课程不存在异常", "通过"],
        ["W06", "任务状态更新", "判定覆盖", "将 TODO 任务改为 COMPLETED", "状态为 COMPLETED 且完成时间非空", "通过"],
        ["W07", "逾期任务统计", "条件覆盖", "创建本用户和他人逾期任务", "只统计本用户逾期任务", "通过"],
        ["W08", "学习时长统计", "语句覆盖", "同一天新增两条学习记录", "按日期汇总为总分钟数", "通过"],
        ["W09", "预算阈值预警", "边界值分析", "预算100，支出85，阈值80", "状态为 WARNING", "通过"],
        ["W10", "重复预算更新", "回归测试", "同月同分类预算重复保存", "更新原预算，不产生唯一键错误", "通过"],
        ["W11", "同日健康记录更新", "回归测试", "同一天健康记录重复保存", "更新原记录，饮水量为最新值", "通过"],
        ["W12", "习惯连续打卡", "循环路径测试", "连续三天打卡", "连续打卡天数为3", "通过"],
        ["W13", "首页聚合", "集成路径测试", "构造课程、任务、健康、预算、目标数据", "看板返回完整聚合结果", "通过"],
    ]
    add_table(doc, ["编号", "测试对象", "方法", "输入/步骤", "预期结果", "实际结果"], white_rows)
    add_p(doc, "（2）结果分析")
    add_p(doc, "本次白盒测试共运行 16 个 JUnit 测试用例，结果为 16 个通过、0 个失败、0 个跳过。测试覆盖了服务层的主要分支，包括正常业务路径、异常路径、边界路径和用户数据隔离路径。")
    add_p(doc, "测试过程中发现健康记录和预算记录存在重复录入时可能触发唯一约束错误的问题。根据测试结果，已将健康记录新增逻辑改为“同日存在则更新”，将预算新增逻辑改为“同月同分类存在则更新”，并增加对应回归测试，修改后测试全部通过。")

    add_heading(doc, "3.2 实验2测试用例及结果分析", 2)
    add_p(doc, "（1）测试用例")
    black_rows = [
        ["B01", "注册登录", "输入合法用户名、昵称、邮箱和密码后注册，再登录", "进入首页看板", "通过"],
        ["B02", "登录异常", "输入不存在用户或错误密码", "页面提示用户名或密码错误", "通过"],
        ["B03", "首页看板", "登录后访问首页", "展示今日课程、待办、健康、预算、目标和打卡统计", "通过"],
        ["B04", "课程管理", "新增一门课程并刷新课程表", "课程列表显示新增课程", "通过"],
        ["B05", "任务管理", "新增 DDL 任务并点击完成", "任务状态更新为已完成", "通过"],
        ["B06", "日程管理", "新增普通日程", "日程列表展示开始时间和结束时间", "通过"],
        ["B07", "学习分析", "新增学习记录", "学习记录列表更新，图表展示学习时长", "通过"],
        ["B08", "健康记录", "新增当天健康记录，再次保存同一天记录", "系统更新原记录而不是报错", "通过"],
        ["B09", "消费预算", "新增支出和预算，支出达到阈值", "预算状态显示 WARNING 或 EXCEEDED", "通过"],
        ["B10", "目标习惯", "新增目标和习惯，点击打卡", "打卡统计增加，习惯列表保持可用", "通过"],
        ["B11", "个人中心", "修改昵称和邮箱", "个人资料保存成功", "通过"],
        ["B12", "未登录访问", "清空本地 token 后访问首页", "自动跳转登录页", "通过"],
    ]
    add_table(doc, ["编号", "测试功能", "测试步骤", "预期结果", "实际结果"], black_rows)
    add_p(doc, "（2）结果分析")
    add_p(doc, "黑盒测试覆盖了用户从注册登录到使用各模块的主要业务流程。测试结果表明系统能够完成大学生个人助手的核心功能，前端页面能够调用后端接口完成数据新增、查询、删除、状态更新和统计展示。")
    add_p(doc, "在黑盒测试中重点关注了重复录入场景和未登录访问场景。重复健康记录和重复预算原先容易造成数据库唯一约束错误，修复后系统表现为更新原有记录，符合用户直觉；未登录访问能够被路由守卫拦截并跳转登录页。")

    add_heading(doc, "4. 实验总结")
    add_p(doc, "", first_line=False)

    doc.add_page_break()
    add_heading(doc, "成绩评定")
    add_table(doc, ["考核标准", "得分"], [
        ["（1）能够按时出勤，不迟到、不早退，实验过程中，具有严谨的学习态度和认真、踏实、一丝不苟的科学作风（5%）", ""],
        ["（2）测试计划全面合理（实验1和实验2各占10%）（20%）", ""],
        ["（3）测试用例设计合理、正确（实验1和实验2各占15%）（30%）", ""],
        ["（4）测试结果分析充分、合理，能体现创新性（实验1和实验2各占10%）（20%）", ""],
        ["（5）待测应用程序编码规范，实验报告规范，不与他人雷同（待测程序占15%，实验报告占10%）（25%）", ""],
    ])
    doc.save(output_path)


if __name__ == "__main__":
    import sys
    build_report(sys.argv[1])

